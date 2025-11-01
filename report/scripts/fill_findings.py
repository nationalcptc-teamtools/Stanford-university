import argparse
import hashlib
import json
import os
import sys
import warnings

import docx
import pandas as pd
import requests
from docx import oxml
from docx.shared import RGBColor

with warnings.catch_warnings():
    # This library is a bit old and warns about pkg_resources in newer Python versions
    warnings.simplefilter("ignore", UserWarning)
    from docxcompose.composer import Composer

from config import Chart, Columns, Labels, Placeholders

MERGED_OUTPUT_DOC = "merged_findings.docx"

with open("charts/findings_chart.json") as f:
    CHART_JSON_TEMPLATE = json.load(f)


def validate_finding_data(finding_data: pd.Series) -> list[str]:
    """
    Run various validation checks on a single finding's data before processing it.
    """
    errors = []
    id = finding_data[Columns.Id]

    # Validate label
    label = finding_data[Columns.Label]
    if label not in Labels.names():
        errors.append(f"Finding {id} has invalid label '{label}'")

    # Validate score
    score = finding_data[Columns.Score]
    if not (0 < score <= 10):
        errors.append(f"Finding {id} has invalid score '{score}'")

    return errors


def validate_findings(df: pd.DataFrame) -> None:
    """
    Run various validation checks on the findings data and stop before processing the findings.
    """
    errors = []

    # Validate columns we require
    required_columns = set(Columns.all())
    missing_columns = required_columns - set(df.columns)
    if missing_columns:
        errors.append(f"Findings sheet is missing required columns: {missing_columns}")

    # Ensure we have no duplicate IDs
    dup_ids = df[df[Columns.Id].duplicated(keep=False)]
    if not dup_ids.empty:
        errors.append(f"Findings sheet has duplicate finding IDs: {dup_ids[Columns.Id].tolist()}")

    # Validate each finding's data
    for _, row_data in df.iterrows():
        finding_errors = validate_finding_data(row_data)
        errors.extend(finding_errors)

    return errors


def sort_findings(df: pd.DataFrame) -> pd.DataFrame:
    """Sort the findings by label severity and then score."""
    label_rank = {label.name: i for i, label in enumerate(Labels)}
    return df.sort_values(
        by=[Columns.Label, Columns.Score],
        key=lambda col: (
            # Sort score by number
            col
            if col.name == Columns.Score
            # Sort label by highest severity to lowest
            else col.map(label_rank)
        ),
        ascending=[True, False],
    )


def read_findings(findings_sheet: str) -> pd.DataFrame:
    """
    Read the findings sheet into a dataframe and sorts the findings. Adds an extra column for each
    finding's label index. Exits if the sheet is not found or the data has validation errors.
    """
    if not os.path.exists(findings_sheet):
        print(f"Findings sheet '{findings_sheet}' not found.")
        sys.exit(1)

    df = pd.read_excel(findings_sheet)

    # Remove any findings with no title, as they are likely not filled out at all
    df = df.dropna(subset=[Columns.Title])

    # Validate the findings
    errors = validate_findings(df)
    if errors:
        print("Errors found in findings sheet:")
        print("\n".join(errors))
        sys.exit(1)

    df = sort_findings(df)

    # Asssign a label index to each finding based on its score
    df = df.assign(_label_index=df.groupby(Columns.Label).cumcount() + 1)
    return df


def set_cell_bg_color(cell: docx.table._Cell, hex_color: str) -> None:
    """
    Set a table cell's background color. The python-docx API does not directly support this, so we
    manipulate the XML directly.
    """
    hex_color = hex_color.lstrip("#").upper()
    shading = oxml.parse_xml(f'<w:shd {oxml.ns.nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)  # noqa: SLF001


def format_run(run: docx.text.run.Run, finding_data: pd.Series) -> None:
    """
    Format a run of text (the base element in a document in the document) with the correct finding
    data. We use placeholders to indicate where to put the data.
    """

    # Color the overall score text to match the label
    if Placeholders.Score.value in run.text:
        label = finding_data[Columns.Label]
        text_color = Labels[label].main_color
        run.font.color.rgb = RGBColor.from_string(text_color.replace("#", ""))

    for placeholder in Placeholders:
        if placeholder.value not in run.text:
            continue

        if placeholder == Placeholders.Index:
            value = finding_data["_label_index"]
        else:
            value = finding_data[Columns.from_placeholder(placeholder)]

        if isinstance(value, float) and value.is_integer():
            # Don't show decimal places for whole numbers
            value = int(value)

        run.text = run.text.replace(placeholder.value, str(value))


def format_row(row: docx.table._Row, finding_data: pd.Series) -> None:
    """Format a table row in the document with the finding data."""

    # Set the left sidebar color
    label = Labels[finding_data[Columns.Label]]
    set_cell_bg_color(row.cells[0], label.sidebar_color)

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                format_run(run, finding_data)


def get_quickchart_chart(chart_str: str) -> bytes:
    """
    Get a chart image from QuickChart given the chart JSON string. Raises an error if the request
    fails.
    """
    resp = requests.post(Chart.QuickChartApiUrl, json={"chart": chart_str})
    resp.raise_for_status()
    return resp.content


def get_chart_image_bytes(finding_data: pd.Series) -> bytes:
    """Generate a chart image from the finding data using QuickChart."""
    chart_json = CHART_JSON_TEMPLATE.copy()
    dataset_json = chart_json["data"]["datasets"][0]

    # Fill in the data points
    columns = finding_data.index.tolist()
    data_start = columns.index(Columns.Severity)
    data_end = columns.index(Columns.Score)
    dataset_json["data"] = finding_data[columns[data_start:data_end]].tolist()

    # Set the colors based on the label
    label = Labels[finding_data[Columns.Label]]
    dataset_json["borderColor"] = label.main_color
    dataset_json["backgroundColor"] = label.background_color

    chart_str = json.dumps(chart_json)
    return get_quickchart_chart(chart_str)


def replace_chart(doc: docx.Document, finding_data: pd.Series) -> None:
    """Replace the default chart image in the document with one generated from the finding data."""
    for rel in doc.part.rels.values():
        if (
            "image" in rel.reltype
            # Only replace the default chart image
            and hashlib.sha256(rel.target_part.blob).hexdigest() == Chart.DefaultImgHash
        ):
            # We access the protected member as the python-docx API does not directly support this
            new_chart_bytes = get_chart_image_bytes(finding_data)
            if new_chart_bytes:
                rel.target_part._blob = new_chart_bytes  # noqa: SLF001


def fill_finding_doc(doc_path: str, finding_data: pd.Series) -> str:
    """Fill a findings document with its data."""
    finding_doc = docx.Document(doc_path)

    for table in finding_doc.tables:
        for row in table.rows:
            format_row(row, finding_data)

    replace_chart(finding_doc, finding_data)

    output_doc = f"{finding_data[Columns.Id]}_filled.docx"
    finding_doc.save(output_doc)

    print(f"Filled out document saved as '{output_doc}'")
    return output_doc


def merge_findings(doc_paths: list[str], output_path: str) -> None:
    """Merge the finding documents into a single document."""

    if not doc_paths:
        print("No documents to merge.")
        return

    print(f"\nMerging {len(doc_paths)} documents into '{output_path}'...")
    base = docx.Document(doc_paths[0])
    base.add_page_break()
    composer = Composer(base)

    for path in doc_paths[1:]:
        try:
            finding = docx.Document(path)
            finding.add_page_break()
            composer.append(finding)
        except Exception as e:
            print(f"Unexpected error merging '{path}', skipping...\n\tError: {e}")

    print("Merge complete")
    composer.save(output_path)


def main(args: argparse.Namespace) -> None:
    findings_sheet = args.findings_sheet
    df = read_findings(findings_sheet)

    selected_ids = set(args.finding_ids or [])
    if selected_ids:
        # If specific IDs are specified by the user, remove any IDs from the dataframe not specified
        missing = selected_ids - set(df[Columns.Id])
        if missing:
            print(f"Warning: Could not find findings for IDs: {', '.join(sorted(missing))}")

        df = df[df[Columns.Id].isin(selected_ids)]
        if df.empty:
            print("No findings matched the provided IDs; nothing to process.")
            return

    doc_paths = []

    # Fill out all docs
    for _, row_data in df.iterrows():
        finding_id = row_data[Columns.Id]

        doc_path = os.path.join(args.findings_dir, f"{finding_id}.docx")
        if not os.path.exists(doc_path):
            print(f"'{finding_id}' document not found in '{args.findings_dir}', skipping...")
            continue

        try:
            filled_doc_path = fill_finding_doc(doc_path, row_data)
        except Exception as e:
            print(f"Unexpected error filling out '{finding_id}' doc, skipping...\n\tError: {e}")
            continue

        doc_paths.append(filled_doc_path)

    # Only merge if want to fill out all findings
    if not selected_ids:
        merge_findings(doc_paths, MERGED_OUTPUT_DOC)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description=(
            "Fill out findings documents from the findings sheet and merge then into a single "
            "document."
        )
    )
    parser.add_argument("findings_sheet", help="Path to the findings sheet")
    parser.add_argument(
        "-i",
        "--finding-id",
        dest="finding_ids",
        action="append",
        metavar="ID",
        help="Only process the specific finding. Use multiple times to include more.",
    )
    parser.add_argument(
        "-d",
        "--findings-dir",
        dest="findings_dir",
        help="Directory containing the findings documents (defaults to the current directory)",
        default=".",
    )

    args = parser.parse_args()
    main(args)
