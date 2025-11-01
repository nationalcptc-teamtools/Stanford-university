import argparse
import io
import json

import docx
import pandas as pd
from docx.shared import Inches, Pt

import fill_findings
from config import Columns, Labels, Style

FINDINGS_DOC = "findings_summary.docx"

# The headers for the findings table
TABLE_HEADERS = ["CCRS ID", "CCRS Value", "Title"]

# The widths of the three columns in the findings table, in inches
COLUMN_WIDTHS_INCHES = [0.95, 1.2, 4.5]

# The margins to apply to table cells
CELL_LEFT_RIGHT_MARGINS = Pt(3.5)
CELL_TOP_BOTTOM_MARGINS = Pt(5)

with open("charts/summary_chart.json") as f:
    SUMMARY_CHART_JSON_TEMPLATE = json.load(f)


def get_summary_chart_bytes(label_counts: list[int]) -> bytes:
    """Generate the summary chart image from the label metrics using QuickChart."""
    chart_json = SUMMARY_CHART_JSON_TEMPLATE.copy()
    chart_json["data"]["labels"] = Labels.labels()
    dataset_json = chart_json["data"]["datasets"][0]

    # Set values for each bar
    dataset_json["data"] = label_counts

    # Set colors for each bar, which represents a label
    dataset_json["borderColor"] = Labels.main_colors()
    dataset_json["backgroundColor"] = Labels.background_colors()

    chart_str = json.dumps(chart_json)
    return fill_findings.get_quickchart_chart(chart_str)


def add_summary_chart(doc: docx.Document, label_counts: list[str, int]) -> None:
    """Add the summary chart as an image to the document."""
    chart_bytes = get_summary_chart_bytes(label_counts)
    doc.add_picture(
        io.BytesIO(chart_bytes),
        width=Inches(Style.SummaryChartWidth),
        height=Inches(Style.SummaryChartHeight),
    )


def set_default_style(doc: docx.Document) -> None:
    """Set the default style for the document based on what is set in the config."""
    font = doc.styles["Normal"].font
    font.name = Style.Font
    font.size = Pt(Style.FontSize)


def set_cell_margins(cell: docx.table._Cell) -> None:
    """
    Apply margins to a table cell via their paragraph indents. Only supports cells with one
    paragraph.
    """
    if not cell.paragraphs:
        return

    p_fmt = cell.paragraphs[0].paragraph_format
    p_fmt.left_indent = CELL_LEFT_RIGHT_MARGINS
    p_fmt.right_indent = CELL_LEFT_RIGHT_MARGINS
    p_fmt.space_before = CELL_TOP_BOTTOM_MARGINS
    p_fmt.space_after = CELL_TOP_BOTTOM_MARGINS


def add_findings_table(doc: docx.Document, findings: pd.DataFrame) -> None:
    """Add a table of findings to the document. Also handles formatting the table."""
    table = doc.add_table(rows=1, cols=3)

    # Format the table
    table.style = "Table Grid"
    table.allow_autofit = False

    # Set column widths
    for idx, col in enumerate(table.columns):
        col.width = Inches(COLUMN_WIDTHS_INCHES[idx])

    # Add and format the header row
    hdr_cells = table.rows[0].cells
    for i, cell in enumerate(hdr_cells):
        run = cell.paragraphs[0].add_run(TABLE_HEADERS[i])
        run.bold = True
        set_cell_margins(cell)

    # Add the findings rows
    for i, finding in findings.iterrows():
        row_cells = table.add_row().cells

        label = finding[Columns.Label]
        label_color = Labels[label].main_color

        # Set the CCRS ID with the text color corresponding to the label
        ccrs_id = f"{finding[Columns.Label]}.{finding['_label_index']}"
        ccrs_run = row_cells[0].paragraphs[0].add_run(ccrs_id)
        ccrs_run.font.color.rgb = docx.shared.RGBColor.from_string(label_color.lstrip("#"))

        # Set the CCRS Value with the label and score
        ccrs_value = f"{Labels[label].label} ({finding[Columns.Score]})"
        row_cells[1].text = ccrs_value

        # Set the Title
        row_cells[2].text = finding["Title"]

        # Format the cells
        set_cell_margins(row_cells[0])
        set_cell_margins(row_cells[1])
        set_cell_margins(row_cells[2])


def generate_findings_summary(findings: pd.DataFrame, output_file) -> None:
    """Generate the findings summary document, which includes the findings chart and table."""
    print("Generating findings summary...")
    doc = docx.Document()
    set_default_style(doc)

    label_counts = findings[Columns.Label].value_counts().reindex(Labels.names(), fill_value=0)
    add_summary_chart(doc, label_counts.tolist())

    p = doc.add_paragraph("\nFindings Matrix:")
    for run in p.runs:
        run.font.size = Pt(12)

    add_findings_table(doc, findings)

    doc.save(output_file)
    print(f"Findings summary saved to {output_file}")


def main(args: argparse.Namespace) -> None:
    findings_sheet = args.findings_sheet
    df = fill_findings.read_findings(findings_sheet)
    generate_findings_summary(df, "findings_summary.docx")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate a findings summary document from the findings sheet."
    )
    parser.add_argument("findings_sheet", help="Path to the findings sheet")

    args = parser.parse_args()
    main(args)
